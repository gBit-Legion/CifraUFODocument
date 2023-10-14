from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT


class VacationReportGenerator:
    def __init__(self):
        self.doc = Document()

    def set_font_size(self, size):
        self.doc.styles['Normal'].font.size = Pt(size)

    def set_paragraph_spacing(self, before=None, after=None):
        for paragraph in self.doc.paragraphs:
            paragraph.paragraph_format.space_before = Inches(before) if before else None
            paragraph.paragraph_format.space_after = Inches(after) if after else None

    def set_paragraph_alignment(self, alignment):
        if alignment == 'left':
            self.doc.styles['Normal'].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        elif alignment == 'right':
            self.doc.styles['Normal'].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        elif alignment == 'center':
            self.doc.styles['Normal'].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    def set_line_spacing(self, value):
        self.doc.styles['Normal'].paragraph_format.space_between = Pt(value)

    def set_page_margins(self, top=None, bottom=None, left=None, right=None):
        if top is not None:
            self.doc.sections[0].top_margin = Inches(top)
        if bottom is not None:
            self.doc.sections[0].bottom_margin = Inches(bottom)
        if left is not None:
            self.doc.sections[0].left_margin = Inches(left)
        if right is not None:
            self.doc.sections[0].right_margin = Inches(right)

    def add_section(self, title, text):
        self.doc.add_heading(title, level=2)
        paragraph = self.doc.add_paragraph(text)
        run = paragraph.runs[0]
        run.bold = True

    def add_information(self, table_name, rows, columns, table_data_list):
        self.doc.add_heading(table_name, level=2)
        table = self.doc.add_table(rows=rows, cols=columns)
        table.style = 'Table Grid'

        for row, data in table.rows:
            for column in table.columns:
                table.cell(row, column).text = table_data_list(row, column)

    def add_list(self, items, list_name):
        self.doc.add_heading(list_name, level=2)
        for item in items:
            p = self.doc.add_paragraph()
            p.add_run(u'\u2022').bold = True
            p.add_run(' ' + item)

    def save_report(self, filename):
        self.doc.save(filename)