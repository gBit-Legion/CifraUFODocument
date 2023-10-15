from docx import Document
from docx.shared import Pt, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT


class DocumentEditor:
    def __init__(self):
        self.doc = Document()
        self.section = self.doc.sections[0]
        self.paragraph = None

    def execute(self, filename, code_table, medium_table, large_table):
        self.set_margins(2, 1, 2, 2)
        self.add_paragraph(left_indent=Inches(4.5), font_size=Pt(10),
                           text='Типовая межотраслевая форма № М-11\nУтверждена распоряжением ОАО "Никелин"\nот 15.12.2008 № 2688р')
        self.add_centered_paragraph('ТРЕБОВАНИЕ-НАКЛАДНАЯ № _______________', is_centered=True, font_size=Pt(12))

        table = self.doc.add_table(rows=4, cols=1, style='Table Grid')
        table_width = Inches(1)
        for row in table.rows:
            for cell in row.cells:
                cell.width = table_width
        page_width = self.section.page_width - self.section.left_margin - self.section.right_margin
        table_indent = page_width - table_width
        table.alignment = 2
        table.alignment_indent = Inches(table_indent)
        table.cell(0, 0).text = code_table[0][0]
        table.cell(0, 0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        table.cell(1, 0).text = code_table[0][1]
        table.cell(1, 0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        table.cell(2, 0).text = code_table[0][2]
        table.cell(2, 0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        table.cell(3, 0).text = code_table[0][3]
        table.cell(3, 0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        self.add_centered_paragraph(
            'Организация_____________________________ОАО "Никелин"_____________________________', is_centered=False,
            font_size=Pt(12))
        self.add_centered_paragraph(
            'Структурное\nподразделение__________________________________________________________________________',
            is_centered=False, font_size=Pt(12))

        table = self.add_table(rows=3, cols=9, style='Table Grid')

        table.cell(0, 0).merge(table.cell(1, 0))
        table.cell(0, 1).merge(table.cell(1, 1))
        table.cell(0, 2).merge(table.cell(0, 3))
        table.cell(0, 4).merge(table.cell(0, 5))
        table.cell(0, 6).merge(table.cell(0, 7))
        table.cell(0, 8).merge(table.cell(1, 8))

        table.cell(0, 0).text = medium_table[0][0]
        table.cell(0, 0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        table.cell(0, 1).text = medium_table[1][0]
        table.cell(0, 1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        table.cell(0, 2).text = medium_table[3][0]
        table.cell(0, 2).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        table.cell(0, 4).text = medium_table[5][0]
        table.cell(0, 4).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        table.cell(0, 6).text = medium_table[7][0]
        table.cell(0, 6).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        table.cell(0, 8).text = medium_table[2][0]
        table.cell(0, 8).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(1, 2).text = medium_table[3][1]
        table.cell(1, 2).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        table.cell(1, 3).text = medium_table[4][0]
        table.cell(1, 3).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        table.cell(1, 4).text = medium_table[5][1]
        table.cell(1, 4).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        table.cell(1, 5).text = medium_table[6][0]
        table.cell(1, 5).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        table.cell(1, 6).text = medium_table[7][1]
        table.cell(1, 6).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        table.cell(1, 7).text = medium_table[8][0]
        table.cell(1, 7).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(2, 0).text = medium_table[0][1]
        table.cell(2, 0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        table.cell(2, 1).text = medium_table[1][1]
        table.cell(2, 1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        table.cell(2, 2).text = medium_table[3][2]
        table.cell(2, 2).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        table.cell(2, 3).text = medium_table[4][1]
        table.cell(2, 3).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        table.cell(2, 4).text = medium_table[5][2]
        table.cell(2, 4).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        table.cell(2, 5).text = medium_table[6][1]
        table.cell(2, 5).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        table.cell(2, 6).text = medium_table[7][2]
        table.cell(2, 6).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        table.cell(2, 7).text = medium_table[8][1]
        table.cell(2, 7).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        table.cell(2, 8).text = medium_table[2][1]
        table.cell(2, 8).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        self.add_centered_paragraph('\nЧерез кого____________________________________________________________________',
                                    is_centered=False, font_size=Pt(12))
        self.add_centered_paragraph(
            '\nЗатребовал_________________________________________Разрешил___________________________________________________',
            is_centered=False, font_size=Pt(12))

        table = self.add_table(rows=len(large_table[9]) + 1, cols=16, style='Table Grid')

        table.cell(0, 0).merge(table.cell(1, 0))
        table.cell(0, 1).merge(table.cell(0, 2))
        table.cell(0, 3).merge(table.cell(1, 3))
        table.cell(0, 4).merge(table.cell(1, 4))
        table.cell(0, 5).merge(table.cell(1, 5))
        table.cell(0, 6).merge(table.cell(1, 6))
        table.cell(0, 7).merge(table.cell(0, 8))
        table.cell(0, 9).merge(table.cell(0, 10))
        table.cell(0, 11).merge(table.cell(1, 11))
        table.cell(0, 12).merge(table.cell(1, 12))
        table.cell(0, 13).merge(table.cell(1, 13))
        table.cell(0, 14).merge(table.cell(1, 14))
        table.cell(0, 15).merge(table.cell(1, 15))

        table.cell(0, 0).text = large_table[0][0]
        table.cell(0, 0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        table.cell(0, 1).text = large_table[9][0]
        table.cell(0, 1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        table.cell(0, 3).text = large_table[1][0]
        table.cell(0, 3).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        table.cell(0, 4).text = large_table[2][0]
        table.cell(0, 4).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        table.cell(0, 5).text = large_table[3][0]
        table.cell(0, 5).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        table.cell(0, 6).text = large_table[14][0]
        table.cell(0, 6).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        table.cell(0, 7).text = large_table[15][0]
        table.cell(0, 7).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        table.cell(0, 9).text = large_table[12][0]
        table.cell(0, 9).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        table.cell(0, 11).text = large_table[4][0]
        table.cell(0, 11).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        table.cell(0, 12).text = large_table[5][0]
        table.cell(0, 12).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        table.cell(0, 13).text = large_table[6][0]
        table.cell(0, 13).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        table.cell(0, 14).text = large_table[7][0]
        table.cell(0, 14).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        table.cell(0, 15).text = large_table[8][0]
        table.cell(0, 15).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(1, 1).text = large_table[9][1]
        table.cell(1, 1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        table.cell(1, 2).text = large_table[10][0]
        table.cell(1, 2).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        table.cell(1, 7).text = 'код'
        table.cell(1, 7).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        table.cell(1, 8).text = large_table[11][0]
        table.cell(1, 8).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        table.cell(1, 9).text = large_table[12][1]
        table.cell(1, 9).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        table.cell(1, 10).text = large_table[13][0]
        table.cell(1, 10).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        for i in range(16):
            table.cell(2, i).text = str(i + 1)
            table.cell(2, i).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        for i, item in enumerate(large_table[0]):
            if i < 1: continue
            table.cell(i + 2, 0).text = item
            paragraph = table.cell(i + 2, 0).paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for i, item in enumerate(large_table[9]):
            if i <= 1: continue
            table.cell(i + 1, 1).text = item
            paragraph = table.cell(i + 1, 1).paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for i, item in enumerate(large_table[10]):
            if i < 1: continue
            table.cell(i + 2, 2).text = item
            paragraph = table.cell(i + 2, 2).paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for i, item in enumerate(large_table[1]):
            if i < 1: continue
            table.cell(i + 2, 3).text = item
            paragraph = table.cell(i + 2, 3).paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for i, item in enumerate(large_table[2]):
            if i < 1: continue
            table.cell(i + 2, 4).text = item
            paragraph = table.cell(i + 2, 4).paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for i, item in enumerate(large_table[3]):
            if i < 1: continue
            table.cell(i + 2, 5).text = item
            paragraph = table.cell(i + 2, 5).paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for i, item in enumerate(large_table[14]):
            if i < 1: continue
            table.cell(i + 2, 6).text = item
            paragraph = table.cell(i + 2, 6).paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for i, item in enumerate(large_table[15]):
            if i <= 1: continue
            table.cell(i + 1, 7).text = item
            paragraph = table.cell(i + 1, 7).paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for i, item in enumerate(large_table[11]):
            if i < 1: continue
            table.cell(i + 2, 8).text = item
            paragraph = table.cell(i + 2, 8).paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for i, item in enumerate(large_table[12]):
            if i <= 1: continue
            table.cell(i + 1, 9).text = item
            paragraph = table.cell(i + 1, 9).paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for i, item in enumerate(large_table[13]):
            if i < 1: continue
            table.cell(i + 2, 10).text = item
            paragraph = table.cell(i + 2, 10).paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for i, item in enumerate(large_table[4]):
            if i < 1: continue
            table.cell(i + 2, 11).text = item
            paragraph = table.cell(i + 2, 11).paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for i, item in enumerate(large_table[5]):
            if i < 1: continue
            table.cell(i + 2, 12).text = item
            paragraph = table.cell(i + 2, 12).paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for i, item in enumerate(large_table[6]):
            if i < 1: continue
            table.cell(i + 2, 13).text = item
            paragraph = table.cell(i + 2, 13).paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for i, item in enumerate(large_table[7]):
            if i < 1: continue
            table.cell(i + 2, 14).text = item
            paragraph = table.cell(i + 2, 14).paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for i, item in enumerate(large_table[8]):
            if i < 1: continue
            table.cell(i + 2, 15).text = item
            paragraph = table.cell(i + 2, 15).paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        self.add_centered_paragraph(
            '\nОтпустил   ____________   ____________   _________________   Получил   ____________   ____________   __________________',
            is_centered=False, font_size=Pt(12))
        self.add_centered_paragraph(
            '\t                (должность)           (подпись)       (расшифровка подписи)\t                      (должность)          (подпись)       (расшифровка подписи)',
            is_centered=False, font_size=Pt(8))
        self.add_centered_paragraph('\nДокумента материала:', is_centered=False, font_size=Pt(12))
        self.add_centered_paragraph('Бухгалтерский документ:', is_centered=False, font_size=Pt(12))

        self.doc.save(filename)

    def set_margins(self, left, right, top, bottom):
        self.section.left_margin = Cm(left)
        self.section.right_margin = Cm(right)
        self.section.top_margin = Cm(top)
        self.section.bottom_margin = Cm(bottom)

    def add_paragraph(self, text=None, left_indent=None, font_size=None):
        self.paragraph = self.doc.add_paragraph(text)
        if left_indent is not None:
            self.paragraph.paragraph_format.left_indent = left_indent
        if font_size is not None:
            self.paragraph.runs[0].font.size = font_size

    def add_centered_paragraph(self, text, is_centered, font_size):
        paragraph = self.doc.add_paragraph()
        if is_centered:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        else:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = paragraph.add_run(text)
        run.font.size = font_size

    def add_table(self, rows, cols, style):
        table = self.doc.add_table(rows=rows, cols=cols, style=style)
        table_width = Inches(1)
        for row in table.rows:
            for cell in row.cells:
                cell.width = table_width
        page_width = self.section.page_width - self.section.left_margin - self.section.right_margin
        table_indent = page_width - table_width
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.alignment_indent = Inches(table_indent)
        return table
